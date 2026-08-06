from __future__ import annotations

import asyncio
import hashlib
import re
from collections.abc import Sequence
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Protocol
from uuid import UUID, uuid4

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from agent.application.models import DependencyUnavailable, InvalidRequest
from agent.settings import Settings


class IngestionRepository(Protocol):
    async def save_document(self, **kwargs: object) -> tuple[int, str]: ...


class ObjectStore(Protocol):
    async def put(self, key: str, data: bytes, content_type: str) -> None: ...

    async def delete(self, key: str) -> None: ...


class DocumentEmbeddingClient(Protocol):
    @property
    def enabled(self) -> bool: ...

    async def embed_documents(self, texts: Sequence[str]) -> list[list[float]]: ...


@dataclass(frozen=True, slots=True)
class IngestionResult:
    document_id: str
    version_number: int
    chunk_count: int
    index_mode: str
    object_key: str


class DocumentIngestionService:
    def __init__(
        self,
        repository: IngestionRepository,
        object_store: ObjectStore,
        embedding: DocumentEmbeddingClient,
        settings: Settings,
    ) -> None:
        self._repository = repository
        self._object_store = object_store
        self._embedding = embedding
        self._settings = settings

    async def ingest(
        self,
        *,
        tenant_id: str,
        knowledge_base_id: str,
        document_id: str | None,
        title: str,
        source_code: str | None,
        file_name: str,
        content_type: str,
        data: bytes,
    ) -> IngestionResult:
        _validate_uuid(tenant_id, "tenant_id")
        _validate_uuid(knowledge_base_id, "knowledge_base_id")
        if document_id:
            _validate_uuid(document_id, "document_id")
        actual_document_id = document_id or str(uuid4())
        version_id = str(uuid4())
        title = " ".join(title.strip().split())
        if not title or len(title) > 300:
            raise InvalidRequest("document title is empty or too long")
        if not data or len(data) > self._settings.max_upload_bytes:
            raise InvalidRequest("document file is empty or exceeds the upload limit")

        try:
            text = await asyncio.to_thread(parse_document, file_name, data)
        except InvalidRequest:
            raise
        except Exception as exc:
            raise InvalidRequest("document could not be parsed") from exc
        if len(text) > self._settings.max_extracted_characters:
            raise InvalidRequest("extracted document text exceeds the configured limit")
        chunks = split_text(
            text,
            self._settings.chunk_size,
            self._settings.chunk_overlap,
        )
        if not chunks:
            raise InvalidRequest("document contains no extractable text")

        embeddings: list[list[float] | None] = [None] * len(chunks)
        index_mode = "KEYWORD"
        if self._embedding.enabled:
            try:
                generated = await self._embed_in_batches(chunks)
                if any(len(vector) != self._settings.embedding_dimension for vector in generated):
                    raise DependencyUnavailable(
                        "embedding dimension does not match database schema"
                    )
                embeddings = list(generated)
                index_mode = "HYBRID"
            except DependencyUnavailable:
                if not self._settings.allow_extractive_fallback:
                    raise

        digest = hashlib.sha256(data).hexdigest()
        safe_name = Path(file_name).name or "document.bin"
        object_key = f"{tenant_id}/{actual_document_id}/{version_id}/{digest[:16]}-{safe_name}"
        await self._object_store.put(object_key, data, content_type)
        try:
            version_number, _index_version = await self._repository.save_document(
                tenant_id=tenant_id,
                knowledge_base_id=knowledge_base_id,
                document_id=actual_document_id,
                document_version_id=version_id,
                title=title,
                source_code=source_code,
                object_key=object_key,
                file_name=safe_name,
                file_size=len(data),
                sha256=digest,
                chunks=chunks,
                embeddings=embeddings,
            )
        except Exception:
            try:
                await self._object_store.delete(object_key)
            except Exception:
                pass
            raise
        return IngestionResult(
            document_id=actual_document_id,
            version_number=version_number,
            chunk_count=len(chunks),
            index_mode=index_mode,
            object_key=object_key,
        )

    async def _embed_in_batches(self, chunks: Sequence[str]) -> list[list[float]]:
        result: list[list[float]] = []
        for start in range(0, len(chunks), 16):
            result.extend(await self._embedding.embed_documents(chunks[start : start + 16]))
        if len(result) != len(chunks):
            raise DependencyUnavailable("embedding response count does not match chunks")
        return result


def parse_document(file_name: str, data: bytes) -> str:
    extension = Path(file_name).suffix.casefold()
    if extension in {".txt", ".md", ".markdown", ".csv"}:
        return _decode_text(data)
    if extension == ".pdf":
        reader = PdfReader(BytesIO(data))
        pages = [
            f"[第 {index} 页]\n{page.extract_text() or ''}"
            for index, page in enumerate(reader.pages, start=1)
        ]
        return "\n\n".join(pages)
    if extension == ".docx":
        document = Document(BytesIO(data))
        values = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values.append("\t".join(cell.text.strip() for cell in row.cells))
        return "\n".join(values)
    if extension == ".xlsx":
        workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
        values: list[str] = []
        try:
            for sheet in workbook.worksheets:
                values.append(f"[工作表：{sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    cells = ["" if value is None else str(value) for value in row]
                    if any(cell.strip() for cell in cells):
                        values.append("\t".join(cells))
        finally:
            workbook.close()
        return "\n".join(values)
    if extension in {".doc", ".xls"}:
        raise InvalidRequest("legacy .doc/.xls is not supported; convert it to .docx/.xlsx")
    raise InvalidRequest(f"unsupported document type: {extension or 'unknown'}")


def split_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    if overlap >= chunk_size:
        raise ValueError("chunk overlap must be smaller than chunk size")
    normalized = re.sub(r"[ \t]+", " ", text.replace("\r\n", "\n").replace("\r", "\n"))
    normalized = re.sub(r"\n{3,}", "\n\n", normalized).strip()
    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(start + chunk_size, len(normalized))
        if end < len(normalized):
            search_start = start + chunk_size // 2
            boundaries = [
                normalized.rfind(separator, search_start, end)
                for separator in ("\n", "。", "！", "？", ";", "；", " ")
            ]
            boundary = max(boundaries)
            if boundary > start:
                end = boundary + 1
        chunk = normalized[start:end].strip()
        if chunk and (not chunks or chunk != chunks[-1]):
            chunks.append(chunk)
        if end >= len(normalized):
            break
        start = max(start + 1, end - overlap)
    return chunks


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise InvalidRequest("text document encoding must be UTF-8 or GB18030")


def _validate_uuid(value: str, field: str) -> None:
    try:
        UUID(value)
    except ValueError as exc:
        raise InvalidRequest(f"{field} is not a valid UUID") from exc
